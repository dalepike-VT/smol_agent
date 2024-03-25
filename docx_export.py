```python
from docx import Document

def save_as_docx(minutes, filename):
    doc = Document()

    doc.add_heading('Meeting Minutes', 0)

    doc.add_heading('Abstract Summary', level=1)
    doc.add_paragraph(minutes['abstract_summary'])

    doc.add_heading('Key Points', level=1)
    doc.add_paragraph(minutes['key_points'])

    doc.add_heading('Action Items', level=1)
    doc.add_paragraph(minutes['action_items'])

    doc.add_heading('Sentiment Analysis', level=1)
    doc.add_paragraph(minutes['sentiment'])

    doc.save(filename)
```